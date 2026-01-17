"""
DOCX text extractor using python-docx.

Extracts text directly from DOCX files without OCR.
Also detects embedded images that might contain text.
"""

import logging
from pathlib import Path

from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE as RT

logger = logging.getLogger(__name__)


def extract_docx(file_path: str | Path) -> tuple[str, bool]:
    """
    Extract text from a DOCX file.

    Extracts all text content from paragraphs, tables, headers, and footers.
    Also checks for embedded images that might contain text.

    Args:
        file_path: Path to the DOCX file

    Returns:
        Tuple of (extracted_text, has_images)
        - extracted_text: All text content from the document
        - has_images: True if document contains embedded images

    Example:
        text, has_images = extract_docx("resume.docx")
        if has_images:
            # Consider OCR for images if needed
            pass
    """
    path = Path(file_path)
    logger.info(f"Extracting DOCX: {path.name}")

    doc = Document(path)
    text_parts = []

    # Extract from main document paragraphs
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text:
            text_parts.append(text)

    # Extract from tables
    for table in doc.tables:
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    row_text.append(cell_text)
            if row_text:
                text_parts.append(" | ".join(row_text))

    # Extract from headers and footers
    for section in doc.sections:
        # Header
        if section.header:
            for paragraph in section.header.paragraphs:
                text = paragraph.text.strip()
                if text:
                    text_parts.append(text)

        # Footer
        if section.footer:
            for paragraph in section.footer.paragraphs:
                text = paragraph.text.strip()
                if text:
                    text_parts.append(text)

    # Check for embedded images
    has_images = _check_for_images(doc)

    full_text = "\n".join(text_parts)
    logger.info(
        f"Extracted {len(full_text)} characters from {path.name}, "
        f"has_images: {has_images}"
    )

    return full_text, has_images


def _check_for_images(doc: Document) -> bool:
    """
    Check if document contains embedded images.

    Images in a DOCX might contain text (scanned signatures,
    scanned sections, etc.) that would need OCR.

    Args:
        doc: python-docx Document object

    Returns:
        True if document contains images
    """
    try:
        # Check document part relationships for images
        for rel in doc.part.rels.values():
            if rel.reltype == RT.IMAGE:
                return True

        # Also check inline shapes
        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                if run._element.xpath(".//a:blip"):
                    return True

    except Exception as e:
        logger.warning(f"Error checking for images: {e}")

    return False


def extract_docx_with_structure(file_path: str | Path) -> dict:
    """
    Extract text from DOCX with structural information.

    Preserves document structure including headings, lists, and tables.
    Useful for section parsing.

    Args:
        file_path: Path to the DOCX file

    Returns:
        Dictionary with structured content:
        {
            "paragraphs": [{"text": str, "style": str}, ...],
            "tables": [[[str, ...], ...], ...],
            "headers": [str, ...],
            "footers": [str, ...],
            "has_images": bool
        }
    """
    path = Path(file_path)
    doc = Document(path)

    result = {
        "paragraphs": [],
        "tables": [],
        "headers": [],
        "footers": [],
        "has_images": False,
    }

    # Extract paragraphs with style info
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text:
            result["paragraphs"].append(
                {
                    "text": text,
                    "style": paragraph.style.name if paragraph.style else "Normal",
                }
            )

    # Extract tables
    for table in doc.tables:
        table_data = []
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells]
            table_data.append(row_data)
        result["tables"].append(table_data)

    # Extract headers
    for section in doc.sections:
        if section.header:
            for paragraph in section.header.paragraphs:
                text = paragraph.text.strip()
                if text:
                    result["headers"].append(text)

    # Extract footers
    for section in doc.sections:
        if section.footer:
            for paragraph in section.footer.paragraphs:
                text = paragraph.text.strip()
                if text:
                    result["footers"].append(text)

    # Check for images
    result["has_images"] = _check_for_images(doc)

    return result
